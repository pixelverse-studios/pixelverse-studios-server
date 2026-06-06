from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "press-releases"
OUTPUT_DIR.mkdir(exist_ok=True)

LOGO_PATH = Path("/Users/phil/PVS-local/Projects/domani/domani-app/assets/AppIcon-sage-512x512.png")
OUTPUT_PATH = OUTPUT_DIR / "domani-launch-press-release.docx"


def set_run(run, size=None, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(document, text="", size=11, bold=False, italic=False, align=None, spacing_after=8):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return paragraph


def main():
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    if LOGO_PATH.exists():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.paragraph_format.space_after = Pt(12)
        logo.add_run().add_picture(str(LOGO_PATH), width=Inches(1.2))

    add_paragraph(
        document,
        "FOR IMMEDIATE RELEASE",
        size=10,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=14,
    )

    heading = add_paragraph(
        document,
        "PixelVerse Studios Launches Domani, a Planning App Designed to Help People Plan Tomorrow, Tonight",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=10,
    )
    heading.paragraph_format.line_spacing = 1.0

    subheading = add_paragraph(
        document,
        "Now available on iOS and Google Play, Domani helps users reduce overwhelm and start each day with clearer priorities",
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=16,
    )

    dateline = document.add_paragraph()
    dateline.paragraph_format.space_after = Pt(10)
    run = dateline.add_run("Englewood Cliffs, NJ, May 21, 2026")
    set_run(run, size=11, bold=True)
    run = dateline.add_run(
        " - PixelVerse Studios has announced the launch of Domani, a modern to-do list app and planning app designed to help people plan tomorrow with more intention. Now available on iOS and Google Play, Domani gives busy people a simple way to organize tasks, focus on priorities, and reduce the morning overwhelm that comes from starting the day without a clear plan."
    )
    set_run(run, size=11)

    body_paragraphs = [
        "Domani is built around evening planning: a reminder to pause at night, choose what matters for tomorrow, and create a daily plan before the next day begins. The app includes task management, priority setting, progress tracking, customizable planning preferences, and a focused daily view so users can move from planning into execution without sorting through an endless list.",
        '"We built Domani around a belief that better days often start the night before," said the PixelVerse Studios team. "When people plan while they are calm, they can wake up with more clarity, less second-guessing, and a stronger sense of what deserves their attention."',
        "Domani is designed for people managing busy schedules, shifting priorities, and the constant pressure to keep track of everything. PixelVerse Studios is encouraging early users to share feedback, with the goal of shaping the app around real routines and everyday planning needs.",
        "Domani is currently available with a 14-day free trial. Early adopters can unlock lifetime access for $9.99, with regular lifetime pricing listed at $34.99. Readers can learn more and find app store links at www.domani-app.com.",
    ]

    for text in body_paragraphs:
        add_paragraph(document, text, size=11, spacing_after=10)

    add_paragraph(document, "About Domani", size=13, bold=True, spacing_after=6)
    add_paragraph(
        document,
        "Domani is a mobile productivity app produced by PixelVerse Studios. Built around the idea of planning tomorrow tonight, Domani helps users reduce overwhelm, organize tasks, focus on priorities, and build a calmer daily planning routine. Domani is available for iOS and Android.",
        size=11,
        spacing_after=16,
    )

    add_paragraph(document, "Media Contact", size=13, bold=True, spacing_after=6)
    contact_lines = [
        "Domani",
        "Website: www.domani-app.com",
        "hello@domani-app.com",
        "Published by PixelVerse Studios",
    ]
    for line in contact_lines:
        add_paragraph(document, line, size=11, spacing_after=2)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
