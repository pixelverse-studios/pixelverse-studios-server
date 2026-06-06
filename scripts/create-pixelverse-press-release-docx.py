from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "press-releases"
OUTPUT_DIR.mkdir(exist_ok=True)

SUPPLIED_LOGO_PATH = Path("/Users/phil/PVS-local/assets/png-file.png")
LOGO_PATH = SUPPLIED_LOGO_PATH if SUPPLIED_LOGO_PATH.exists() else ROOT / "src" / "media" / "logo-email.png"
OUTPUT_PATH = OUTPUT_DIR / "pixelverse-north-new-jersey-chamber-press-release.docx"


def set_run(run, size=None, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(document, text="", size=11, bold=False, italic=False, align=None, spacing_after=8):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return paragraph


def main():
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    if LOGO_PATH.exists():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.paragraph_format.space_after = Pt(12)
        logo.add_run().add_picture(str(LOGO_PATH), width=Inches(2.35))

    add_paragraph(
        document,
        "FOR IMMEDIATE RELEASE",
        size=10,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=14,
    )

    heading = add_paragraph(
        document,
        "PixelVerse Studios Joins the North New Jersey Chamber of Commerce",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=16,
    )
    heading.paragraph_format.line_spacing = 1.0

    dateline = document.add_paragraph()
    dateline.paragraph_format.space_after = Pt(10)
    run = dateline.add_run("Bergen County, NJ, May 20, 2026")
    set_run(run, size=11, bold=True)
    run = dateline.add_run(
        " — PixelVerse Studios has joined the North New Jersey Chamber of Commerce, marking a new step in the company's commitment to supporting local businesses across Bergen County and the tri-state area."
    )
    set_run(run, size=11)

    body_paragraphs = [
        "The team at PixelVerse Studios helps businesses build modern, easy-to-use digital products that generate real results. Its services include custom websites, mobile apps, local SEO, website audits, business automation, and custom internal tools such as dashboards and CRM systems that help companies operate more efficiently.",
        "\"Local businesses are the foundation of strong communities, and we want to help them operate and compete in an increasingly digital world. By joining the North New Jersey Chamber of Commerce, we're excited to build relationships with business owners who care deeply about their work and their communities.\"",
    ]

    for text in body_paragraphs:
        add_paragraph(document, text, size=11, spacing_after=10)

    offer = document.add_paragraph()
    offer.paragraph_format.space_after = Pt(10)
    parts = [
        ("To welcome fellow chamber members, PixelVerse Studios is offering ", False),
        ("25% off setup costs", True),
        (" and an additional ", False),
        ("10% off retainer costs", True),
        (" for businesses that sign on before ", False),
        ("August 1, 2026", True),
        (".", False),
    ]
    for text, bold in parts:
        run = offer.add_run(text)
        set_run(run, size=11, bold=bold)

    add_paragraph(
        document,
        "PixelVerse Studios works with businesses locally across Bergen County and the tri-state area, as well as remote clients across industries.",
        size=11,
        spacing_after=16,
    )

    add_paragraph(document, "About PixelVerse Studios", size=13, bold=True, spacing_after=6)
    add_paragraph(
        document,
        "PixelVerse Studios is a Bergen County-based digital studio helping businesses use technology to grow, operate more efficiently, and compete in an increasingly digital world. The studio serves clients throughout the tri-state area and anywhere digitally.",
        size=11,
        spacing_after=16,
    )

    add_paragraph(document, "Media Contact", size=13, bold=True, spacing_after=6)
    contact_lines = [
        "PixelVerse Studios",
        "Bergen County, NJ",
        "info@pixelversestudios.io",
        "914-297-8215",
        "www.pixelversestudios.io",
    ]
    for line in contact_lines:
        add_paragraph(document, line, size=11, spacing_after=2)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
